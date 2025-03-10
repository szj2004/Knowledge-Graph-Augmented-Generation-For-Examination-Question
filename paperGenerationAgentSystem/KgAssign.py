import ast
import os
import random
import time
from functools import partial
import numpy as np
#
# import pynvml
import sum_score as ss
from generateQuestion import generate
import save_mysql
import gpt




def getKG(knowledge):
    prompt = "\n从以上输入中适当地提取出历史知识点，输出为列表"
    example = "\n 例如：秦始皇统一六国，建立中国历史上第一个封建王朝\n 输出为['秦始皇','统一六国','中国第一个封建王朝','秦朝']" \
              "\n禁止其他任何无意义的输出，回答仅仅是个列表"
    res=gpt.query_chatgpt_model(knowledge+prompt+example)
    return res


def giveReview(ql):
    paper=""
    for t in ql:
        type=t[0]
        ques=t[2]
        if type == 1 or type == 2:
            paper+=ques+"\n"
        else:
            mat = t[4]
            paper+=mat
            paper+=ques+"\n"
    prompt ="\n以上是一份历史试卷的所有题目，分别评价选择题，简答题（如果有），材料题，的难度，知识点，出题质量等等方面，最后给出一个综合评价。"
    res=gpt.query_chatgpt_model(paper+prompt)
    # print(res)
    return res



def KgExtraction(knowledge):


    # 接收用户的输入 知识点用\n分开
    kgs = knowledge.split("\n")

    user_list=[]
    # 让llm提取知识点 kg是用户输入的知识点，kp是llm分割的知识点
    for kg in kgs:
        # 这里k是重要参数，控制着用户输入的知识点在所有知识点中的占比>>>>>>>>>>>>>>>>>>>>>>>
        k=int(len(kg)/50)+1
        kp_list = ast.literal_eval(getKG(kg))
        # 再加入从表中寻找到的相似的知识点
        list1=[]
        for kp in kp_list:
            # 这里的k决定着从一个kp里提取几个相似的知识点
            list1.extend(save_mysql.get_kp_all(kp,k=1))
        print(list1)
        if len(list1)>=k:
            # print(list1)
            #定义排序函数
            def custom_sort_function(item, const):
                return ss.sort_score(kp_list=const,kg=item)

            partial_func = partial(custom_sort_function, const=kp_list)
            sorted_list = sorted(list1, key=partial_func)
            # print(sorted_list)
            user_list.extend(sorted_list[:k])
        else:
            user_list.extend(list1[:])

    return user_list

def KgGenerate(sum,knowledge):
    # print("start of extract:" + str(time.time() ))
    #user表示用户输入的，supply表示我补充的
    user_list=KgExtraction(knowledge)
    # print(user_list)
    # print("\nend of extract:" + str(time.time() ))
    user_num=len(user_list)
    supply_num=max(sum-user_num,0)
    print("\nstart of supply:" + str(time.time()))
    supply_list=save_mysql.get_supply(number=supply_num,pc_list=user_list[:])
    print("\n all the end:"+str(time.time()))
    kg_list=user_list+supply_list
    # print(user_list)
    # print(supply_list)
    # print(kg_list)
    return  user_list,supply_list,kg_list



#设置各个题目类型的数量，当只有题目总数，使用预设

# def QuestionNumAssign(sum):
#     min_choice=int(sum*0.6)
#     max_choice=int(sum*0.7)
#     min_analyse=int(sum*0.2)
#     max_analyse=int(sum*0.3)
#     min_material=int(sum*0.2)
#     max_material=int(sum*0.3)
#     if min_choice == max_choice:
#         choice = min_choice
#     else:
#         choice = random.randint(min_choice, max_choice)
#     remaining = sum - choice
#     if max(min_analyse, remaining - max_material) == min(max_analyse, remaining):
#         analyse= min(max_analyse, remaining)
#     if  max(min_analyse, remaining - max_material) < min(max_analyse, remaining):
#         analyse = random.randint(max(min_analyse, remaining - max_material), min(max_analyse, remaining))
#     else :
#         analyse = random.randint(min(max_analyse, remaining),max(min_analyse, remaining - max_material))
#
#     material = sum - choice - analyse
#     return choice, analyse, material

# 分割大语言模型生成的问题和答案
# text 大语言模型生成的结果
#选择与问答是 res 加类型 返回 ques和ans 材料则多了个mat

def deal(text, type):
    lines = text.split("\n")
    if type == 1:
        state = 1
        content, answer = '', ''
        for (i, line) in enumerate(lines):
            if len(line) == 0:
                continue
            if i == 0 and (line[-1] == '：' or line[-1] == ':'):
                continue
            if '答案' in line[:30]:
                state = 2
            if len(line) == 3:
                continue

            if state == 1:
                content = content + line + '\n'
            elif state == 2:
                answer = answer + line + '\n'
        content = content.strip()
        answer = answer.strip()
        return content[3:], "【答案】"+answer[3:]
    elif type == 2:
        ques,ans="",""
        state=1
        for (i, line) in enumerate(lines):
            line = line.strip()
            if len(line) == 0:
                continue
            elif '答案' in line[:30]:
                state = 2
                ans =  line + '\n'
            elif state == 1:
                ques = ques + line + '\n'
            elif state == 2:
                ans = ans + line + '\n'
        return ques[3:].replace('\n',''),"【参考答案】"+ans[3:-1]
    elif type == 3:
        # print(text)
        mat1, mat2, ques1, ques2, ans1, ans2 = '', '', '', '', '', ''
        state = 1
        for line in lines:
            line=line.replace("*", "")
            line = line.strip()
            if len(line) == 0:
                continue
            elif line[:3] == '材料一':
                mat1 = line
            elif line[:3] == '材料二':
                mat2 = line
                state = 2
            elif line[:3] == '问题一':
                ques1 = line
                state = 3
            elif line[:3] == '问题二':
                ques2 = line
                state = 4
            elif line[:2] == '答案' and ans1 == '':
                ans1 = line
                state = 5
            elif line[:2] == '答案':
                ans2 = line
                state = 6
            else:
                if state == 1:
                    mat1 = mat1 + line
                elif state == 2:
                    mat2 = mat2 + line
                elif state == 3:
                    ques1 = ques1 + line
                elif state == 4:
                    ques2 = ques2 + line
                elif state == 5:
                    ans1 = ans1 + line
                elif state == 6:
                    ans2 = ans2 + line

        mat=mat1+"\n"+mat2
        ques="(1)"+ques1[4:]+"\n(2)"+ques2[4:]
        ans= "(1)【参考答案】"+ans1[3:]+"\n(2)【参考答案】"+ans2[3:]
        return mat,ques,ans
#l,m,h分别代表低，中，高三种难度题目的比例
def ProducePaper(knowledge,choice=2,analyse=2,material=2,l=0.7,m=0.2,h=0.1,title="history"):

    # choice,analyse,material=QuestionNumAssign(sum=sum)
    sum=choice+analyse+material
#计算出各种题目的比例,维护一个队列 type_list 简称tl，分别是choice_l,choice_m,choice_h,analyse_l,analyse_m,analyse_h,material_l,material_m,material_h，这也是遍历顺序
    tl=[0] * 9
    tl[0]=int(l*choice)
    tl[1]=int(m*choice)
    tl[2]=choice-tl[0]-tl[1]
    tl[3]=int(l*analyse)
    tl[4]=int(m*analyse)
    tl[5]=analyse-tl[4]-tl[3]
    tl[6] = int(l * material)
    tl[7] = int(m * material)
    tl[8] = material - tl[6] - tl[7]
    # print(tl)
    if knowledge.endswith("\n"):
        knowledge= knowledge[:-1]
    print(1)
    user_list,supply_list,kg_list=KgGenerate(sum=sum,knowledge=knowledge)
    #index 是知识点的指针
    index=0
    print(2)
    #存储题目 ql = question_list   元素内容为[type method ques ans (mat)]
    ql=[]
    for i, t in enumerate(tl):
        for x in range(0,t):
            print("i:"+str(i))
            temp=[]
            # 123对应类型
            type = int(i / 3) + 1
            temp.append(type)
            #，method不同难度对应的方法
            methods = ['Rand', 'KP', 'KKP', 'KAG']
            method = i % 3
            method = random.randint(method, method + 1)
            method = methods[method]
            temp.append(method)
            res=generate(data=kg_list[index],type=type,method=method)
            index=index+1
            if  type == 3:
                # print(res)
                mat,ques,ans=deal(res, type)
                temp.append(ques)
                temp.append(ans)
                temp.append(mat)
            else :
                ques, ans = deal(res, type)
                temp.append(ques)
                temp.append(ans)
            ql.append(temp)
            print(3)
    generateDocx(title, ql)
    return 0

#生成文档
def generateDocx(title,ql):
    # python操作word
    # 导入库
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    # 新建空白文档
    doc = Document()
    # 添加标题（0相当于文章的题目，默认级别是1，级别范围为0-9）
    heading=doc.add_heading(title, level=1)
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #设置题目
    for i ,t in enumerate(ql):
        type=t[0]
        method=t[1]
        ques=t[2]
        if len(t) == 5:
            mat=t[4]
        if type == 1:
            if i == 0:
                p = doc.add_paragraph()
                run = p.add_run("\n一、选择题,每题只有一个正确选项。")
                run.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(12)
            p = doc.add_paragraph()
            run = p.add_run(str(i+1)+"."+ques)
            run.font.name = '宋体'
            run.font.size = Pt(10)
        elif type == 2:
            if ql[i - 1][0] == 1:
                p = doc.add_paragraph()
                run = p.add_run("\n二、简答题,根据题目简述答案。")
                run.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(12)
            p = doc.add_paragraph()
            run = p.add_run(str(i+1)+"."+ques)
            run.font.name = '宋体'
            run.font.size = Pt(10)

        elif type == 3:
            if ql[i-1][0] == 2:
                p = doc.add_paragraph()
                run = p.add_run("\n三、材料题,根据材料回答问题。")
                run.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(12)

            p = doc.add_paragraph()
            run = p.add_run(str(i+1)+".\n"+mat)
            run.font.name = '宋体'
            run.font.size = Pt(10)

            p = doc.add_paragraph()
            run = p.add_run(ques)
            run.font.name = '宋体'
            run.font.size = Pt(10)
    # 设置答案
    doc.add_page_break()
    heading = doc.add_heading(title+'【参考答案】', level=1)
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    for i ,t in enumerate(ql):
        ans=t[3]
        run = p.add_run(str(i+1)+".\n"+ans+"\n")
        run.font.name = '宋体'
        run.font.size = Pt(10)

    #设置评价
    doc.add_page_break()
    heading = doc.add_heading(title + '【综合评价】', level=1)
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(giveReview(ql))
    run.font.name = '宋体'
    run.font.size = Pt(10)

    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def create_page_number(doc):
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        run = paragraph.add_run()
        # 设置页码
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        # 设置居中对齐
        paragraph.alignment = 1
        # 设置加粗
        run.font.bold = True

    create_page_number(doc)

    # 保存文件（当前目录下）
    doc.save('paper.docx')



def set_gpu():
    import pynvml
    try:
        # 初始化 NVIDIA 管理库
        pynvml.nvmlInit()
        # 获取 GPU 设备数量
        gpu_count = pynvml.nvmlDeviceGetCount()
        min_utlization = 100
        idlist = []
        min_id=1;
        # 遍历所有 GPU 设备
        for gpu_id in range(gpu_count):
            handle = pynvml.nvmlDeviceGetHandleByIndex(gpu_id)
            # 获取 GPU 信息
            gpu_utilization = (int)(pynvml.nvmlDeviceGetUtilizationRates(handle).gpu)
            if(gpu_utilization<min_utlization):
                min_utlization=gpu_utilization
                min_id=gpu_id
            if(gpu_utilization<20):
                idlist.append(gpu_id)
        #设置
        if (len(idlist) == 0):
            idlist.append(min_id)
        os.environ["CUDA_VISIBLE_DEVICES"] = ','.join(map(str, idlist))
        # 释放 NVIDIA 管理库
        pynvml.nvmlShutdown()
        return min_id,idlist
    except pynvml.NVMLError as error:
        print(f"NVIDIA Management Library Error: {error}")


if __name__=="__main__":
    set_gpu()
    # ql=ProducePaper(sum=1,knowledge="商鞅变法\n",choice=0,material=1,analyse=0)
    # print(ql)
    ql=[[1, 'KP', '《史记·商君列传》记载，商鞅是法家的代表人物之一，他推行变法，强化中央集权，制定严苛的法律，强调刑法和奖赏的明确，以达到国家治理的目的。商鞅的政治理念最能体现下列哪一特点？\nA. 强调道德教化，注重礼仪  \nB. 推崇儒家思想，提倡仁政  \nC. 强化法治，推行严格的法律制度  \nD. 提倡分封制，重视地方自治', '【答案】C'], [1, 'KAG', '秦始皇在公元前210年实施的“焚书坑儒”政策的目的主要是？\nA. 压制学术思想，强化中央集权  \nB. 消除对其政权的历史质疑  \nC. 促进科技发展，遏制思想束缚  \nD. 保护传统文化，恢复礼仪制度', '【答案】A'], [2, 'KP', '请简述阿兹特克文明的文化特点。', '【参考答案】（1）阿兹特克文明具有较高的文化水平，表现为先进的农业技术、复杂的社会组织以及丰富的宗教仪式。\n（2）阿兹特克人精通天文学，能够精准地计算日历，建立了复杂的历法系统。\n（3）他们创造了文字系统，用于记录历史和宗教事务。\n（4）阿兹特克人还在艺术和建筑上有显著成就，建造了宏伟的神庙和城市，特别是在首都特诺奇蒂特兰。\n（5）此外，阿兹特克文明注重战争与祭祀，祭祀活动常常涉及人祭，以求取神明的庇佑。\n'], [2, 'KAG', '请简述1978年十一届三中全会对中国社会的影响。', '【参考答案】\n（1）1978年十一届三中全会标志着中国改革开放的起点，邓小平提出的邓小平理论成为改革的指导思想。\n（2）中国造船业在1978年之前较为封闭，改革后逐渐融入市场，并于1982年正式进入世界造船市场，推动了中国工业的现代化。\n（3）文化大革命结束后，人民代表大会制度得到了进一步完善，社会主义民主政治建设在1982年取得新进展，为中国的政治体制改革打下基础。\n'], [3, 'Rand', '(1)\n(2)', '(1)【参考答案】\n(2)【参考答案】', '材料一：在抗日战争时期，许多中国人民和抗日力量采取了与殖民当局不合作的态度，尤其是在一些沿海地区的抗日根据地，民众普遍拒绝与日本占领者进行合作，并积极参与抗日活动。以厦门为例，当地人民不仅不与日本占领当局进行任何形式的经济或文化合作，还积极参与情报传递和游击战斗，形成了强有力的抗日力量。\n材料二：在印度的独立运动中，甘地领导的非暴力不合作运动取得了显著成效。通过号召印度民众拒绝与英国殖民当局的任何合作，包括拒绝英籍商品、停止缴纳税款和抵制英资企业等，印度人民通过不合作的方式表达了反对殖民统治的决心。这种策略不仅引发了大规模的社会动员，也让殖民政府陷入了巨大的经济困境。**问题一**：根据材料一和材料二，分析与殖民当局不合作的共同特点。**答案**：与殖民当局不合作的共同特点包括：一是通过拒绝与殖民政府的经济和文化合作，抵制殖民统治；二是强调民众参与，广泛动员人民力量，形成对抗殖民统治的强大社会力量；三是通过不合作行动造成殖民政府的经济和政治困境，迫使其对抗压力作出让步。**问题二**：根据材料一和材料二，评估“与殖民当局不合作”对抗殖民统治的效果。**答案**：不合作运动通常能有效削弱殖民当局的控制力和经济基础。在中国，尤其是抗日战争期间，民众的反抗行动给日本占领当局造成了相当大的压力，减少了资源供给，并增强了抗日力量的凝聚力。在印度，不合作运动通过广泛的社会动员，打击了英国的经济利益，推动了印度的独立进程。因此，不合作策略能够在一定程度上打乱殖民当局的统治计划，加速殖民地的独立进程。'], [3, 'KKP', '(1)根据材料一、二，分析洋务新政未能达到保护国家安全、抵抗外敌侵略的原因。\n(2)根据材料一，说明洋务新政在引入机器生产技术方面的成效。', '(1)【参考答案】原因在于改革缺乏系统性规划与根本性变革，过度依赖外部技术，未能做到自主创新，且改革力度不足，未能彻底提升清朝的整体国力。\n(2)【参考答案】洋务新政在引入机器生产技术方面取得了一定成效，推动了铁路、电报、船舰等基础设施建设，并促进了部分工业的发展，为国家自卫能力的提升奠定了基础。', '材料一：洋务新政时期（19世纪60-90年代），洋务派主张“自强”与“求富”，提出引进西方的机器生产技术，建立现代化工业。为了提升国家的综合国力，他们推动了铁路、电报、船舰等基础设施建设，并积极引进先进的军事技术，加强国防力量。通过这一系列的改革措施，洋务派希望能够增强清朝的自卫能力，抵御外敌侵略，保护国家的安全。\n材料二：尽管洋务新政在技术引进和工业发展方面取得了一定成效，但由于改革过程中缺乏系统性规划与根本性变革，未能有效改变国家整体弱势。由于对现代化的认识偏差，以及过度依赖西方技术，未能做到真正的自主创新，导致洋务派未能完全达成保护国家安全和抵抗外敌侵略的目标。尽管洋务新政促进了军事力量的提升，但清朝最终未能避免外来侵略，且国内局势依然动荡不安。']]
    generateDocx("2023年普通高等学校招生全国统一考试（全国甲卷）",ql)